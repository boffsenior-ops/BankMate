import docx
import os

os.makedirs("test_docs", exist_ok=True)

doc = docx.Document()
doc.add_heading("BankMate Ipoteka Krediti Yo'riqnomasi", 0)

doc.add_heading("1. Umumiy qoidalar", level=1)
doc.add_paragraph("Ushbu yo'riqnoma BankMate mijozi uchun ipoteka krediti ajratish shartlarini belgilaydi.")

doc.add_heading("2. Foiz stavkasi va shartlari", level=1)
doc.add_paragraph("Ipoteka krediti uchun yillik foiz stavkasi 17.5% qilib belgilangan. Agar mijoz oylik maoshini bizning bank orqali olsa, imtiyozli 16% stavka qo'llaniladi.")
doc.add_paragraph("Maksimal muddat 20 yil (240 oy). Dastlabki badal kamida 15% bo'lishi shart.")

doc.save("test_docs/kredit_test.docx")
print("test_docs/kredit_test.docx created successfully.")
